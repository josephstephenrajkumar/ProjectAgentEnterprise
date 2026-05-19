import docx
from docx.shared import Pt, Inches
import os

doc = docx.Document()
doc.add_heading('Architect Interview Materials', 0)

files = [
    '/home/joseph/.gemini/antigravity/brain/6bcf0be0-bdbb-4fa4-b282-f2522f190408/architect_interview_prep.md',
    '/home/joseph/.gemini/antigravity/brain/6bcf0be0-bdbb-4fa4-b282-f2522f190408/architecture_diagram.md'
]

for file in files:
    try:
        with open(file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        in_code_block = False
        is_mermaid = False
        
        for line in lines:
            stripped = line.strip()
            
            # Handle code blocks
            if stripped.startswith('```mermaid'):
                in_code_block = True
                is_mermaid = True
                
                # Insert the generated diagram image instead of code
                if os.path.exists('architecture.png'):
                    doc.add_picture('architecture.png', width=Inches(6.0))
                continue
                
            elif stripped.startswith('```'):
                in_code_block = not in_code_block
                if not in_code_block:
                    is_mermaid = False
                continue
                
            if in_code_block:
                if not is_mermaid:
                    p = doc.add_paragraph()
                    p.add_run(line.rstrip('\n')).font.name = 'Courier New'
                continue

            if not stripped:
                doc.add_paragraph()
                continue
                
            # Basic Markdown Parsing
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('**') and ':**' in stripped:
                p = doc.add_paragraph()
                split = stripped.split(':**', 1)
                p.add_run(split[0].replace('**', '') + ':').bold = True
                p.add_run(split[1].replace('**', ''))
            elif stripped.startswith('* **'):
                p = doc.add_paragraph(style='List Bullet')
                if ':**' in stripped:
                    split = stripped.split(':**', 1)
                    p.add_run(split[0].replace('* **', '') + ':').bold = True
                    p.add_run(split[1].replace('**', ''))
                elif '***' in stripped or '**' in stripped:
                     p.add_run(stripped.replace('* **', '').replace('**', ''))
                else:
                    p.add_run(stripped.replace('* **', '').replace('**', ''))
            elif stripped.startswith('* '):
                 p = doc.add_paragraph(stripped[2:], style='List Bullet')
            else:
                doc.add_paragraph(stripped.replace('**', ''))
                
        doc.add_page_break()
    except Exception as e:
        print(f"Error reading {file}: {e}")

output_path = '/home/joseph/projectAgent/Architect_Interview_Materials.docx'
doc.save(output_path)
print(f"Successfully generated {output_path}")
